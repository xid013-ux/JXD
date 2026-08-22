# -*- coding: utf-8 -*-
"""《融资 BP 框架（备忘）》—— 一页一页告诉客户团队怎么做。"""
import sys, os
SK = '.claude/skills/gtht-research-report/assets'
sys.path.insert(0, SK)
os.environ['DOCX_OUT'] = 'output/BP框架/融资BP框架（备忘）.docx'
from _helpers import *
from _helpers import _q
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
sys.path.insert(0, 'output/BP框架')
from bp_content import PARTS, EXTRAS, CHECKS, DESENSITIZE, MATERIALS, STEPS

DEEP = RGBColor(0x12, 0x3F, 0x63)
GRAY = RGBColor(0x59, 0x59, 0x59)


def _sp(p, before, after, line=300):
    pPr = p._p.get_or_add_pPr()
    e = OxmlElement('w:spacing')
    e.set(qn('w:before'), str(before)); e.set(qn('w:after'), str(after))
    e.set(qn('w:line'), str(line)); e.set(qn('w:lineRule'), 'auto')
    pPr.append(e)
    return pPr


def _ind(pPr, left=0, hanging=0, first=0):
    e = OxmlElement('w:ind')
    e.set(qn('w:left'), str(left))
    if hanging: e.set(qn('w:hanging'), str(hanging))
    else: e.set(qn('w:firstLine'), str(first))
    pPr.append(e)


def subtitle(t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(p, 0, 240)
    r = p.add_run(t)
    rf = r._element.get_or_add_rPr().get_or_add_rFonts()
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia'): rf.set(qn(a), '宋体')
    rf.set(qn('w:hint'), 'eastAsia')
    r.font.size = Pt(10.5); r.font.color.rgb = GRAY
    return p


def plain(t, size=12):
    """正文段落，首行缩进两字。"""
    p = doc.add_paragraph()
    pPr = _sp(p, 60, 60, 330)
    _ind(pPr, first=480)
    r = p.add_run(_q(t))
    rf = r._element.get_or_add_rPr().get_or_add_rFonts()
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia'): rf.set(qn(a), '宋体')
    rf.set(qn('w:hint'), 'eastAsia')
    r.font.size = Pt(size)
    return p


def lead(t):
    """部分标题下面的一行小字。"""
    p = doc.add_paragraph()
    pPr = _sp(p, 40, 80, 300)
    _ind(pPr, left=0)
    r = p.add_run(_q('这页说明白：' + t))
    rf = r._element.get_or_add_rPr().get_or_add_rFonts()
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia'): rf.set(qn(a), '宋体')
    rf.set(qn('w:hint'), 'eastAsia')
    r.font.size = Pt(10.5); r.font.color.rgb = GRAY
    return p


def label(t):
    """小标题：先准备 / 这页这样写 / 提醒。"""
    p = doc.add_paragraph()
    pPr = _sp(p, 120, 40, 300)
    pPr.append(OxmlElement('w:keepNext'))
    _ind(pPr, left=0)
    r = p.add_run(t)
    rf = r._element.get_or_add_rPr().get_or_add_rFonts()
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia'): rf.set(qn(a), '黑体')
    rf.set(qn('w:hint'), 'eastAsia')
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = DEEP
    return p


def bullet(t):
    p = doc.add_paragraph()
    pPr = _sp(p, 20, 20, 300)
    _ind(pPr, left=640, hanging=220)
    r = p.add_run('· ' + _q(t))
    rf = r._element.get_or_add_rPr().get_or_add_rFonts()
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia'): rf.set(qn(a), '宋体')
    rf.set(qn('w:hint'), 'eastAsia')
    r.font.size = Pt(11.5)
    return p


def warn(t):
    p = doc.add_paragraph()
    pPr = _sp(p, 20, 140, 300)
    _ind(pPr, left=640, hanging=0, first=0)
    r = p.add_run(_q(t))
    rf = r._element.get_or_add_rPr().get_or_add_rFonts()
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia'): rf.set(qn(a), '宋体')
    rf.set(qn('w:hint'), 'eastAsia')
    r.font.size = Pt(11.5)
    return p


def part_head(no, name, pages):
    p = doc.add_paragraph(style='Heading 3')
    pPr = p._p.get_or_add_pPr()
    npr = OxmlElement('w:numPr')
    il = OxmlElement('w:ilvl'); il.set(qn('w:val'), '0'); npr.append(il)
    ni = OxmlElement('w:numId'); ni.set(qn('w:val'), '0'); npr.append(ni)
    pPr.append(npr)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '280'); sp.set(qn('w:after'), '60'); pPr.append(sp)
    pPr.append(OxmlElement('w:keepNext'))
    r = p.add_run('%s ｜ %s' % (no, name))
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:hint'), 'eastAsia')
    r.font.size = Pt(13)
    r2 = p.add_run('    %s' % pages)
    rf = r2._element.get_or_add_rPr().get_or_add_rFonts()
    rf.set(qn('w:eastAsia'), '宋体'); rf.set(qn('w:hint'), 'eastAsia')
    r2.bold = False; r2.font.size = Pt(10.5); r2.font.color.rgb = GRAY
    return p


# ==================================================================
title('融资 BP 框架（备忘）')
subtitle('一页一页告诉团队：每部分准备什么材料、写什么内容、画什么图')

h2('先说清楚')

plain('这份东西不是范文，是骨架。它把一份融资 BP 拆成 17 页，'
      '每一页告诉你三件事：先准备哪些材料、这一页怎么写、最容易在哪里翻车。'
      '内容各家自己填，同样的骨架不同公司做出来完全不一样，这是正常的。')
plain('两条通用规矩，贯穿全篇。第一，每一页的标题写一句结论，不要写名词——'
      '把"产品优势"改成"良率从 65% 提到 92%，2025 年 11 月起批量交付"，'
      '页面内容就是这句话的证据。第二，能画图就别写字，能列表就别写段落，'
      '一页只讲一件事。')
plain('正式版做 15—25 页。做完之后删成 10—12 页当路演版，再压成一页纸发给'
      '第一次接触的机构。三个版本共用一套数据，只是详略不同。')

table(headers=['页', '部分', '建议页数', '页', '部分', '建议页数'],
      rows=[[PARTS[i][0], PARTS[i][1], PARTS[i][2],
             PARTS[i + 9][0] if i + 9 < len(PARTS) else '',
             PARTS[i + 9][1] if i + 9 < len(PARTS) else '',
             PARTS[i + 9][2] if i + 9 < len(PARTS) else '']
            for i in range(9)],
      widths=[6, 22, 12, 6, 22, 12],
      caption='表1  17 页一览',
      aligns=['center', 'left', 'center', 'center', 'left', 'center'])

h2('一页一页怎么做')

for no, name, pages, goal, prep, how, tip in PARTS:
    part_head(no, name, pages)
    lead(goal)
    label('先准备这些材料')
    for x in prep: bullet(x)
    label('这一页这样写')
    for x in how: bullet(x)
    label('提醒')
    warn(tip)

h2('有这些情况，再加几页')

plain('下面七页不是每家都要写。有对应的事实就单独做一页，放在正文里比塞进附录管用。')

table(headers=['加什么页', '什么情况下加', '这页怎么做'],
      rows=[list(e) for e in EXTRAS],
      widths=[14, 20, 66],
      caption='表2  加分页',
      aligns=['left', 'left', 'left'])

h2('怎么推进')

plain('材料不齐就开始做页面，只能写形容词。先按下表把料收齐，缺什么落实到人。')

table(headers=['类别', '要收的材料', '找谁要'],
      rows=[list(m) for m in MATERIALS],
      widths=[12, 62, 26],
      caption='表3  备料清单',
      aligns=['center', 'left', 'center'])

plain('料齐了按下面的顺序做。先写字后排版，是为了别把时间花在美化上而内容没想清楚。')

table(headers=['顺序', '做什么', '大概用时'],
      rows=[list(s) for s in STEPS],
      widths=[10, 72, 18],
      caption='表4  做的顺序',
      aligns=['center', 'left', 'center'])

h2('交之前对一遍')

label('这几条逐条打勾')
for i, c in enumerate(CHECKS, 1):
    bullet('%s' % c)

label('发给外部之前，记得脱密')
for d in DESENSITIZE:
    bullet(d)
warn('脱密只做遮盖，不改数字。脱密版和完整版的数字必须一样。')

import _fixup
