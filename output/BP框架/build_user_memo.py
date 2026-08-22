# -*- coding: utf-8 -*-
"""《燕京中发融资 BP 框架编写指引》—— 按部门内部报告体例排版。
内容取自用户修订版 PPT，逐字保留。"""
import sys, os
SK = '.claude/skills/gtht-research-report/assets'
sys.path.insert(0, SK)
os.environ['DOCX_OUT'] = 'output/BP框架/燕京中发融资BP框架编写指引.docx'
from _helpers import *
from _helpers import _q
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
sys.path.insert(0, 'output/BP框架')
from zf_user_content import COVER, OUTLINE, CHAPTERS

DEEP = RGBColor(0x12, 0x3F, 0x63)
GRAY = RGBColor(0x59, 0x59, 0x59)
CN_IDX = '一二三四五六七八九十'


def _sp(p, before, after, line=330):
    pPr = p._p.get_or_add_pPr()
    e = OxmlElement('w:spacing')
    e.set(qn('w:before'), str(before)); e.set(qn('w:after'), str(after))
    e.set(qn('w:line'), str(line)); e.set(qn('w:lineRule'), 'auto')
    pPr.append(e)
    return pPr


def _ind(pPr, left=0, hanging=0, first=0):
    e = OxmlElement('w:ind')
    e.set(qn('w:left'), str(left))
    if hanging: e.set(qn('w:hanging'), str(hanging))
    elif first: e.set(qn('w:firstLine'), str(first))
    pPr.append(e)


def _run(p, t, name='宋体', size=12, color=None, bold=False):
    r = p.add_run(_q(t))
    rf = r._element.get_or_add_rPr().get_or_add_rFonts()
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia'): rf.set(qn(a), name)
    rf.set(qn('w:hint'), 'eastAsia')
    r.font.size = Pt(size)
    if color is not None: r.font.color.rgb = color
    r.bold = bold
    return r


def subtitle(t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(p, 0, 240)
    _run(p, t, size=10.5, color=GRAY)


def body(t):
    """正文段落：宋体小四、首行缩进两字。"""
    p = doc.add_paragraph(); _ind(_sp(p, 60, 60), first=480)
    _run(p, t)


def note_line(t):
    """页面说明行：小一号、灰色、不缩进。"""
    p = doc.add_paragraph(); _ind(_sp(p, 40, 60, 300), left=0)
    _run(p, t, size=10.5, color=GRAY)


def sub_label(t):
    """段内小标题：黑体、深蓝、不入目录。"""
    p = doc.add_paragraph()
    pPr = _sp(p, 120, 40, 300); pPr.append(OxmlElement('w:keepNext')); _ind(pPr, left=0)
    _run(p, t, name='黑体', size=11, color=DEEP, bold=True)


def item(i, t):
    """编号条目：（1）……，悬挂缩进对齐。"""
    p = doc.add_paragraph(); _ind(_sp(p, 20, 20, 310), left=880, hanging=480)
    _run(p, '（%d）%s' % (i, t), size=11.5)


def plain_item(t):
    p = doc.add_paragraph(); _ind(_sp(p, 20, 60, 310), left=880)
    _run(p, t, size=11.5)


def page_head(num, name, no):
    p = doc.add_paragraph(style='Heading 3')
    pPr = p._p.get_or_add_pPr()
    npr = OxmlElement('w:numPr')
    il = OxmlElement('w:ilvl'); il.set(qn('w:val'), '0'); npr.append(il)
    ni = OxmlElement('w:numId'); ni.set(qn('w:val'), '0'); npr.append(ni)
    pPr.append(npr)
    sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '260'); sp.set(qn('w:after'), '60')
    pPr.append(sp); pPr.append(OxmlElement('w:keepNext'))
    r = p.add_run('%s  %s' % (num, name))
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:hint'), 'eastAsia')
    r.font.size = Pt(12)
    r2 = p.add_run('　（BP 第 %s 页）' % no)
    rf = r2._element.get_or_add_rPr().get_or_add_rFonts()
    rf.set(qn('w:eastAsia'), '宋体'); rf.set(qn('w:hint'), 'eastAsia')
    r2.bold = False; r2.font.size = Pt(10.5); r2.font.color.rgb = GRAY
    # 目录里只显示编号与名称
    HEADINGS[-1] = (3, '%s  %s' % (num, name), HEADINGS[-1][2])
    return p


# ==================================================================
title('燕京中发融资 BP 框架编写指引')
subtitle(COVER['sub'])

h2('全篇大纲概览')
body('本指引与《燕京中发　融资 BP 框架》教学材料逐页对应，'
     '用于公司各部门在编制融资商业计划书过程中对照使用。'
     '全篇分五章、共 25 页，每一页列明三项内容：编写前需要准备的材料、'
     '本页的写法，以及编写时应当注意的事项。')
body('章节顺序建议不做调整。' + OUTLINE['lead'] +
     '本轮融资为混合所有制改革引入战略投资者，方案须报国资监管部门审核并在'
     '北京产权交易所公开挂牌，材料在表述上应当与审核口径保持一致，'
     '涉及价格、比例、时间等尚未确定的事项一律表述为拟定安排。')

rows = []
for cno, cname, cgoal, pages in CHAPTERS:
    for i, (no, name, lead, _p, _h, _t, _ht) in enumerate(pages):
        rows.append([cname if i == 0 else '', no, name, lead])
table(headers=['章', '页', '页面名称', '本页说明'],
      rows=rows, widths=[11, 6, 26, 57],
      caption='表1  全篇 25 页构成',
      aligns=['center', 'center', 'left', 'left'])

for ci, (cno, cname, cgoal, pages) in enumerate(CHAPTERS, start=2):
    h2(cname)
    if cgoal:
        body('本章对应 BP %s，共 %d 页。%s' % (cno, len(pages), cgoal))
    else:
        body('本章对应 BP %s，共 %d 页，为第 %s 页至第 %s 页。'
             % (cno, len(pages), pages[0][0], pages[-1][0]))
    for pi, (no, name, lead, prep, how, tip, how_title) in enumerate(pages, start=1):
        page_head('%d.%d' % (ci, pi), name, no)
        note_line('这页说清楚：' + lead)
        if prep:
            sub_label('先准备这些材料')
            for i, x in enumerate(prep, 1): item(i, x)
        sub_label(how_title)
        for i, x in enumerate(how, 1): item(i, x)
        sub_label('提醒')
        plain_item(tip)

make_toc(doc)
import _fixup
