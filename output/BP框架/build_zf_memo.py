# -*- coding: utf-8 -*-
"""《燕京中发融资 BP 框架（备忘）》"""
import sys, os
SK = '.claude/skills/gtht-research-report/assets'
sys.path.insert(0, SK)
os.environ['DOCX_OUT'] = 'output/BP框架/燕京中发融资BP框架（备忘）.docx'
from _helpers import *
from _helpers import _q
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
sys.path.insert(0, 'output/BP框架')
from zf_content import BRIEF, CHAPTERS, SPECIALS, MATERIALS, STEPS, CHECKS

DEEP = RGBColor(0x12, 0x3F, 0x63)
GRAY = RGBColor(0x59, 0x59, 0x59)


def _sp(p, before, after, line=300):
    pPr = p._p.get_or_add_pPr()
    e = OxmlElement('w:spacing')
    e.set(qn('w:before'), str(before)); e.set(qn('w:after'), str(after))
    e.set(qn('w:line'), str(line)); e.set(qn('w:lineRule'), 'auto')
    pPr.append(e)
    return pPr


def _ind(pPr, left=0, hanging=0, first=0):
    e = OxmlElement('w:ind')
    e.set(qn('w:left'), str(left))
    if hanging: e.set(qn('w:hanging'), str(hanging))
    elif first: e.set(qn('w:firstLine'), str(first))
    pPr.append(e)


def _font(r, name='宋体', size=12, color=None, bold=False):
    rf = r._element.get_or_add_rPr().get_or_add_rFonts()
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia'): rf.set(qn(a), name)
    rf.set(qn('w:hint'), 'eastAsia')
    r.font.size = Pt(size)
    if color is not None: r.font.color.rgb = color
    r.bold = bold
    return r


def subtitle(t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(p, 0, 240)
    _font(p.add_run(t), size=10.5, color=GRAY)


def plain(t):
    p = doc.add_paragraph()
    pPr = _sp(p, 60, 60, 330); _ind(pPr, first=480)
    _font(p.add_run(_q(t)), size=12)


def lead(t, prefix='这页说清楚：'):
    p = doc.add_paragraph()
    pPr = _sp(p, 40, 80, 300); _ind(pPr, left=0)
    _font(p.add_run(_q(prefix + t)), size=10.5, color=GRAY)


def label(t):
    p = doc.add_paragraph()
    pPr = _sp(p, 120, 40, 300); pPr.append(OxmlElement('w:keepNext')); _ind(pPr, left=0)
    _font(p.add_run(t), name='黑体', size=11, color=DEEP, bold=True)


def bullet(t):
    p = doc.add_paragraph()
    pPr = _sp(p, 20, 20, 300); _ind(pPr, left=640, hanging=220)
    _font(p.add_run('· ' + _q(t)), size=11.5)


def warn(t):
    p = doc.add_paragraph()
    pPr = _sp(p, 20, 140, 300); _ind(pPr, left=640)
    _font(p.add_run(_q(t)), size=11.5)


def page_head(no, name):
    p = doc.add_paragraph(style='Heading 3')
    pPr = p._p.get_or_add_pPr()
    npr = OxmlElement('w:numPr')
    il = OxmlElement('w:ilvl'); il.set(qn('w:val'), '0'); npr.append(il)
    ni = OxmlElement('w:numId'); ni.set(qn('w:val'), '0'); npr.append(ni)
    pPr.append(npr)
    sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '280'); sp.set(qn('w:after'), '60')
    pPr.append(sp); pPr.append(OxmlElement('w:keepNext'))
    r = p.add_run('%s ｜ %s' % (no, name))
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:hint'), 'eastAsia')
    r.font.size = Pt(13)


# ==================================================================
title('燕京中发融资 BP 框架（备忘）')
subtitle('混合所有制改革引入战略投资者 · 材料编写指引')

h2('先说清楚')
for k, v in BRIEF:
    label(k); plain(v)

plain('两条通用规矩，贯穿全篇。第一，每页标题写一句结论，不写名词——'
      '把"公司优势"改成"蓝帽子批文 12 个，国产同类企业中数量第一"，'
      '页面内容就是这句话的证据。第二，能画图就别写字，能列表就别写段落，'
      '一页只讲一件事。')

rows = []
for cno, cname, cgoal, pages in CHAPTERS:
    for i, (no, name, goal, _p, _h, _t) in enumerate(pages):
        rows.append([cname if i == 0 else '', no, name, goal.replace('这页说清楚：', '')])
table(headers=['章', '页', '页面名称', '这页说清楚什么'],
      rows=rows, widths=[12, 6, 24, 58],
      caption='表1  全篇 25 页一览',
      aligns=['center', 'center', 'left', 'left'])

for cno, cname, cgoal, pages in CHAPTERS:
    h2(cname)
    lead(cgoal, prefix='对应 BP %s，共 %d 页。' % (cno, len(pages)))
    for no, name, goal, prep, how, tip in pages:
        page_head(no, name)
        lead(goal.replace('这页说清楚：', ''))
        label('先准备这些材料')
        for x in prep: bullet(x)
        label('这一页这样写')
        for x in how: bullet(x)
        label('提醒')
        warn(tip)

h2('这个项目特别要注意的六件事')
plain('下面几条是燕京中发这个项目独有的，普通融资 BP 上没有，写之前先记住。')
table(headers=['事项', '怎么处理'],
      rows=[list(x) for x in SPECIALS], widths=[16, 84],
      caption='表2  特别注意事项', aligns=['center', 'left'])

h2('动手之前，先把料收齐')
plain('料不齐就开始做页面，只能写形容词。按下表收料，缺什么落实到人。')
table(headers=['类别', '要收的材料', '找谁要'],
      rows=[list(m) for m in MATERIALS], widths=[12, 62, 26],
      caption='表3  备料清单', aligns=['center', 'left', 'center'])
plain('料齐了按下面的顺序做。先写字后排版，是为了别把时间花在美化上而内容没想清楚。')
table(headers=['顺序', '做什么', '大概用时'],
      rows=[list(s) for s in STEPS], widths=[10, 72, 18],
      caption='表4  做的顺序', aligns=['center', 'left', 'center'])

h2('交之前对一遍')
for c in CHECKS: bullet(c)

import _fixup
