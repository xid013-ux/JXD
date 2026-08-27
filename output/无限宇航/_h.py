# -*- coding: utf-8 -*-
"""按《融资参考案例0408》体例构建 Word 的辅助函数。
排版参数全部取自该模版：正文宋体/Times New Roman 小四、首行缩进2字符、1.5倍行距；
表格100%宽度居中、外框0.75磅内线0.25磅、表内五号、表头加粗居中。"""
import copy
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TPL = '_模板底稿.docx'

def new_doc():
    d = Document(TPL)
    body = d.element.body
    for el in list(body):
        if not el.tag.endswith('}sectPr'):
            body.remove(el)
    return d


def _keepnext(p):
    """在 pStyle 之后插入 keepNext（CT_PPr 要求 pStyle 排第一）"""
    pPr = p._element.get_or_add_pPr()
    if pPr.find(qn('w:keepNext')) is not None:
        return
    kn = OxmlElement('w:keepNext')
    ps = pPr.find(qn('w:pStyle'))
    if ps is None:
        pPr.insert(0, kn)
    else:
        ps.addnext(kn)

def _run(p, text, bold=False, sz=None):
    r = p.add_run(text)
    r.bold = bold
    if sz:
        r.font.size = Pt(sz)
    rPr = r._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.insert(0, rf)
    rf.set(qn('w:hint'), 'eastAsia')
    return r

def title(d, text):
    p = d.add_paragraph(style='Heading 1')
    # 模版标题样式自带前置分页，首页不需要
    pPr = p._element.get_or_add_pPr()
    pb = pPr.find(qn('w:pageBreakBefore'))
    if pb is None:
        pb = OxmlElement('w:pageBreakBefore'); pPr.append(pb)
    pb.set(qn('w:val'), '0')
    _run(p, text)
    return p

def h(d, text):
    """四级标题：1、xxx"""
    p = d.add_paragraph(style='004四级标题')
    _run(p, text, bold=True)
    return p

def para(d, text):
    p = d.add_paragraph(style='005正文')
    _run(p, text)
    return p

def unit(d, text):
    """表格前单位，右对齐五号"""
    p = d.add_paragraph(style='表格前单位')
    _run(p, text, sz=10.5)
    return p

def note(d, text):
    """数据来源注：七号居中（部门体例）"""
    p = d.add_paragraph(style='表格后说明')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, text, sz=5.5)
    return p

def cap(d, text):
    """表/图标题：五号加粗居中，编号与题名之间用空格"""
    p = d.add_paragraph(style='表格后说明')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _keepnext(p)
    _run(p, text, bold=True, sz=10.5)
    return p

def figure(d, path, caption, source):
    """插图：宽度与正文栏宽匹配，图题在下方"""
    from PIL import Image
    from docx.shared import Emu
    W = 5074920
    w, h = Image.open(path).size
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _keepnext(p)
    p.add_run().add_picture(path, width=Emu(W), height=Emu(int(W*h/w)))
    cap(d, caption)
    note(d, source)

def _cell(tc, text, bold, align, first_row, last_row, width):
    """按 CT_TcPr 规定顺序重建单元格属性：tcW -> tcBorders -> vAlign"""
    tcPr = tc._element.get_or_add_tcPr()
    for child in list(tcPr):
        tcPr.remove(child)
    cw = OxmlElement('w:tcW'); cw.set(qn('w:w'), str(width))
    cw.set(qn('w:type'), 'pct'); tcPr.append(cw)
    if first_row or last_row:
        b = OxmlElement('w:tcBorders')
        if first_row:
            t = OxmlElement('w:top'); t.set(qn('w:val'),'single'); t.set(qn('w:sz'),'12')
            t.set(qn('w:space'),'0'); t.set(qn('w:color'),'010000'); b.append(t)
        bo = OxmlElement('w:bottom'); bo.set(qn('w:val'),'single')
        bo.set(qn('w:sz'), '12' if last_row else '4')
        bo.set(qn('w:space'),'0'); bo.set(qn('w:color'),'010000'); b.append(bo)
        tcPr.append(b)
    va = OxmlElement('w:vAlign'); va.set(qn('w:val'), 'center'); tcPr.append(va)
    p = tc.paragraphs[0]
    _keepnext(p)
    p.alignment = {'c': WD_ALIGN_PARAGRAPH.CENTER,
                   'l': WD_ALIGN_PARAGRAPH.LEFT,
                   'r': WD_ALIGN_PARAGRAPH.RIGHT}[align]
    if text:
        _run(p, text, bold=bold, sz=10.5)

def table(d, headers, rows, widths=None, aligns=None):
    n = len(headers)
    t = d.add_table(rows=0, cols=n)
    tbl = t._element
    tblPr = tbl.tblPr
    for tag in ('w:tblW','w:jc','w:tblBorders','w:tblLook','w:tblStyle'):
        e = tblPr.find(qn(tag))
        if e is not None:
            tblPr.remove(e)
    w = OxmlElement('w:tblW'); w.set(qn('w:w'),'5000'); w.set(qn('w:type'),'pct'); tblPr.append(w)
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'),'center'); tblPr.append(jc)
    bs = OxmlElement('w:tblBorders')
    for side, sz in (('top','12'),('left','12'),('bottom','12'),('right','12'),
                     ('insideH','4'),('insideV','4')):
        e = OxmlElement('w:'+side); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),sz)
        e.set(qn('w:space'),'0'); e.set(qn('w:color'),'010000'); bs.append(e)
    tblPr.append(bs)
    lk = OxmlElement('w:tblLook')
    for k,v in (('w:val','04A0'),('w:firstRow','1'),('w:lastRow','0'),
                ('w:firstColumn','1'),('w:lastColumn','0'),
                ('w:noHBand','0'),('w:noVBand','1')):
        lk.set(qn(k), v)
    tblPr.append(lk)

    if widths is None:
        widths = [round(5000/n)] * n
    else:
        s = sum(widths); widths = [round(w*5000/s) for w in widths]
    if aligns is None:
        aligns = ['c'] * n

    all_rows = [headers] + rows
    for ri, rd in enumerate(all_rows):
        row = t.add_row()
        trPr = row._element.get_or_add_trPr()
        if ri == 0:
            th = OxmlElement('w:tblHeader'); trPr.append(th)
        ka = OxmlElement('w:cantSplit'); trPr.insert(0, ka)
        for ci, val in enumerate(rd):
            _cell(row.cells[ci], val, bold=(ri == 0),
                  align=('c' if ri == 0 else aligns[ci]),
                  first_row=(ri == 0), last_row=(ri == len(all_rows)-1),
                  width=widths[ci])
    return t

def save(d, path):
    cp = d.core_properties
    cp.author = ''; cp.last_modified_by = ''; cp.title = ''
    cp.comments = ''; cp.category = ''; cp.keywords = ''; cp.subject = ''
    d.save(path)
    print('saved', path)
