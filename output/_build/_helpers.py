# -*- coding: utf-8 -*-
import copy, os
from docx import Document
from docx.shared import Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

TPL = '/tmp/claude-0/-home-user-JXD/b33c1a84-553f-5c48-a1f8-b4ca1cd8f155/scratchpad/template.docx'
import os
OUT = os.environ.get('DOCX_OUT','/home/user/JXD/output/out.docx')
CH  = '/home/user/JXD/output/charts/'

doc = Document(TPL)
body = doc.element.body
sectPr = body.find(qn('w:sectPr'))
for child in list(body):
    if child is not sectPr:
        body.remove(child)

def _rpr(el, bold=False, sz=None, color='000000'):
    rPr = OxmlElement('w:rPr')
    f = OxmlElement('w:rFonts'); f.set(qn('w:hint'), 'eastAsia'); rPr.append(f)
    if bold: rPr.append(OxmlElement('w:b'))
    c = OxmlElement('w:color'); c.set(qn('w:val'), color); rPr.append(c)
    if sz:
        s = OxmlElement('w:sz'); s.set(qn('w:val'), str(sz)); rPr.append(s)
    el.append(rPr)

def title(text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    for tag, attrs in (('w:keepNext', {}), ('w:widowControl', {}),
                       ('w:spacing', {'w:before':'156','w:after':'156','w:line':'360','w:lineRule':'auto'}),
                       ('w:jc', {'w:val':'center'})):
        e = OxmlElement(tag)
        for k, v in attrs.items(): e.set(qn(k), v)
        pPr.append(e)
    r = p.add_run(text)
    rf = r._element.get_or_add_rPr().get_or_add_rFonts()
    for a in ('w:ascii','w:hAnsi','w:cs','w:eastAsia'): rf.set(qn(a), '宋体')
    rf.set(qn('w:hint'), 'eastAsia')
    r.bold = True; r.font.size = Pt(18)
    return p

HEADINGS = []          # (level, display_text, bookmark_name)
_BK = [1000]
CN_NUM = '一二三四五六七八九十'

def _bookmark(p, name):
    from docx.oxml import OxmlElement
    bs = OxmlElement('w:bookmarkStart'); bs.set(qn('w:id'), str(_BK[0])); bs.set(qn('w:name'), name)
    be = OxmlElement('w:bookmarkEnd'); be.set(qn('w:id'), str(_BK[0]))
    _BK[0] += 1
    pPr = p._p.find(qn('w:pPr'))
    p._p.insert(list(p._p).index(pPr) + 1 if pPr is not None else 0, bs)
    p._p.append(be)

def h2(text):
    p = doc.add_paragraph(style='Heading 2')
    pPr = p._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing'); sp.set(qn('w:before'),'156'); sp.set(qn('w:after'),'156'); pPr.append(sp)
    r = p.add_run(text); r.bold = True
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:hint'),'eastAsia')
    n = sum(1 for lv, _, _ in HEADINGS if lv == 2)
    name = '_Toc90%04d' % len(HEADINGS)
    HEADINGS.append((2, CN_NUM[n] + '、 ' + text, name))
    _bookmark(p, name)
    return p

def h3(text):
    p = doc.add_paragraph(style='Heading 3')
    pPr = p._p.get_or_add_pPr()
    npr = OxmlElement('w:numPr')
    il = OxmlElement('w:ilvl'); il.set(qn('w:val'),'0'); npr.append(il)
    ni = OxmlElement('w:numId'); ni.set(qn('w:val'),'0'); npr.append(ni)
    pPr.append(npr)
    sp = OxmlElement('w:spacing'); sp.set(qn('w:before'),'156'); sp.set(qn('w:after'),'156'); pPr.append(sp)
    r = p.add_run(text)
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:hint'),'eastAsia')
    name = '_Toc90%04d' % len(HEADINGS)
    HEADINGS.append((3, text, name))
    _bookmark(p, name)
    return p

def para(text):
    p = doc.add_paragraph(style='Normal Indent')
    r = p.add_run(text)
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:hint'),'eastAsia')
    return p

def note(text):
    p = doc.add_paragraph(style='表格后说明')
    r = p.add_run(text)
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:hint'),'eastAsia')
    return p

def blank():
    return doc.add_paragraph()

CONTENT_EMU = 5074920

def figure(fname, caption, source):
    img = Image.open(CH + fname); w, h = img.size
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    kn = OxmlElement('w:keepNext'); pPr.append(kn)
    p.add_run().add_picture(CH + fname, width=Emu(CONTENT_EMU), height=Emu(int(CONTENT_EMU * h / w)))
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(caption); r.bold = True; r.font.size = Pt(10.5)
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:hint'),'eastAsia')
    note(source)

def set_borders(el, top=None, bottom=None, left=None, right=None, iH=None, iV=None):
    b = OxmlElement('w:tblBorders') if el.tag == qn('w:tblPr') else OxmlElement('w:tcBorders')
    for tag, sz in (('w:top',top),('w:left',left),('w:bottom',bottom),('w:right',right),
                    ('w:insideH',iH),('w:insideV',iV)):
        if sz is None: continue
        e = OxmlElement(tag); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),str(sz))
        e.set(qn('w:space'),'0'); e.set(qn('w:color'),'010000'); b.append(e)
    el.append(b)

def table(headers, rows, widths, caption=None, source=None, aligns=None):
    """widths: list of pct-of-table numbers summing to 5000-ish scale handled internally."""
    if caption:
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = cp._p.get_or_add_pPr(); pPr.append(OxmlElement('w:keepNext'))
        r = cp.add_run(caption); r.bold = True; r.font.size = Pt(10.5)
        r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:hint'),'eastAsia')
    ncol = len(headers)
    tot = sum(widths)
    pcts = [int(round(w / tot * 5000)) for w in widths]
    tbl = doc.add_table(rows=0, cols=ncol)
    tblPr = tbl._tbl.tblPr
    for ch in list(tblPr):
        if ch.tag in (qn('w:tblW'), qn('w:tblBorders'), qn('w:jc'), qn('w:tblStyle')):
            tblPr.remove(ch)
    tw = OxmlElement('w:tblW'); tw.set(qn('w:w'),'5352'); tw.set(qn('w:type'),'pct'); tblPr.append(tw)
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'),'center'); tblPr.append(jc)
    set_borders(tblPr, top=12, bottom=12, left=12, right=12, iH=4, iV=4)

    def fill(cells, texts, header):
        for i, (c, t) in enumerate(zip(cells, texts)):
            tcPr = c._tc.get_or_add_tcPr()
            for ch in list(tcPr):
                if ch.tag == qn('w:tcW'): tcPr.remove(ch)
            w = OxmlElement('w:tcW'); w.set(qn('w:w'),str(pcts[i])); w.set(qn('w:type'),'pct'); tcPr.append(w)
            if header: set_borders(tcPr, top=12, bottom=4)
            else: set_borders(tcPr, top=4)
            va = OxmlElement('w:vAlign'); va.set(qn('w:val'),'center'); tcPr.append(va)
            c.paragraphs[0].text = ''
            lines = str(t).split('\n')
            for j, ln in enumerate(lines):
                p = c.paragraphs[0] if j == 0 else c.add_paragraph()
                pPr = p._p.get_or_add_pPr()
                if header: pPr.append(OxmlElement('w:keepNext'))
                a = OxmlElement('w:jc')
                if header: a.set(qn('w:val'),'center')
                else:
                    al = (aligns[i] if aligns else 'left')
                    a.set(qn('w:val'), al)
                pPr.append(a)
                r = p.add_run(ln)
                rPr = r._element.get_or_add_rPr()
                rPr.get_or_add_rFonts().set(qn('w:hint'),'eastAsia')
                if header: r.bold = True
                r.font.size = Pt(10.5)
                r.font.color.rgb = RGBColor(0,0,0)

    hr = tbl.add_row()
    trPr = hr._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:cantSplit')); trPr.append(OxmlElement('w:tblHeader'))
    j = OxmlElement('w:jc'); j.set(qn('w:val'),'center'); trPr.append(j)
    fill(hr.cells, headers, True)
    for row in rows:
        tr = tbl.add_row()
        trPr = tr._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
        j = OxmlElement('w:jc'); j.set(qn('w:val'),'center'); trPr.append(j)
        fill(tr.cells, row, False)
    if source: note(source)
    return tbl


# ---------------- 目录（与模板一致的 Word TOC 域） ----------------
def _w(t): return int(t)

def _cjk_units(s):
    return sum(1.0 if ord(c) > 0x2E80 else 0.5 for c in s)

PAGE_H = 16834 - 1440 - 1440          # 可用高度（缇）
TEXT_W = 11909 - 1800 - 1800          # 可用宽度（缇）

def _est_body_pages(doc):
    """粗略估算每个书签所在页码；Word 打开时会通过 updateFields 自动刷新为准确值。"""
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    pages = {}
    page, used = 1, 0
    def advance(hgt):
        nonlocal page, used
        if used + hgt > PAGE_H:
            page += 1; used = hgt
        else:
            used += hgt
    for ch in doc.element.body.iterchildren():
        if ch.tag == qn('w:p'):
            p = Paragraph(ch, doc)
            for bs in ch.findall(qn('w:bookmarkStart')):
                pages[bs.get(qn('w:name'))] = page
            st = p.style.name if p.style is not None else 'Normal'
            txt = p.text
            if ch.findall('.//' + qn('w:drawing')):
                ext = ch.find('.//' + qn('wp:extent'))
                advance(int(ext.get('cy')) // 635 + 120 if ext is not None else 2500)
                continue
            if st in ('heading 2', 'Heading 2'):
                advance(420 + 312)
            elif st in ('heading 3', 'Heading 3'):
                advance(360 + 312)
            elif st == 'Normal Indent':
                advance(max(1, -(-int(_cjk_units(txt) + 2) // 34)) * 360 + 156)
            elif st == '表格后说明':
                advance(max(1, -(-int(_cjk_units(txt)) // 40)) * 260 + 60)
            else:
                advance(max(1, -(-int(_cjk_units(txt)) // 38)) * 260 if txt.strip() else 240)
        elif ch.tag == qn('w:tbl'):
            t = Table(ch, doc)
            grid = [int(g.get(qn('w:w'))) for g in ch.findall(qn('w:tblGrid') + '/' + qn('w:gridCol'))]
            tot = sum(grid) or 5000
            for row in t.rows:
                hh = 0
                for i, c in enumerate(row.cells):
                    colw = TEXT_W * 1.07 * (grid[i] / tot) if i < len(grid) else TEXT_W / len(row.cells)
                    per = max(4, int(colw // 210))
                    hh = max(hh, max(1, -(-int(_cjk_units(c.text)) // per)) * 252)
                advance(hh + 120)
    return pages

def make_toc(doc):
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    body = doc.element.body
    pages = _est_body_pages(doc)
    # 目录自身占用的页数
    toc_h = 620 + len(HEADINGS) * 400
    toc_pages = max(1, -(-toc_h // PAGE_H))
    offset = toc_pages

    def esc(s):
        return (s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))

    NS = nsdecls('w', 'w14')
    parts = []
    parts.append(
        '<w:p %s><w:pPr><w:pStyle w:val="TOC"/><w:spacing w:before="156" w:after="156" '
        'w:line="360" w:lineRule="auto"/><w:rPr><w:rFonts w:eastAsia="宋体"/><w:b/>'
        '<w:color w:val="000000"/><w:sz w:val="28"/><w:szCs w:val="28"/></w:rPr></w:pPr>'
        '<w:r><w:rPr><w:rFonts w:eastAsia="宋体" w:hint="eastAsia"/><w:b/><w:color w:val="000000"/>'
        '<w:sz w:val="28"/><w:szCs w:val="28"/></w:rPr><w:t>目录</w:t></w:r></w:p>' % NS)

    for i, (lvl, text, name) in enumerate(HEADINGS):
        pg = pages.get(name, 1) + offset
        first = (i == 0)
        if lvl == 2:
            ppr = ('<w:pPr><w:pStyle w:val="TOC1"/><w:tabs><w:tab w:val="right" w:leader="dot" '
                   'w:pos="8299"/></w:tabs><w:spacing w:line="360" w:lineRule="auto"/>'
                   '<w:rPr><w:noProof/><w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr></w:pPr>')
            rs = '<w:rStyle w:val="a5"/><w:bCs w:val="0"/><w:noProof/><w:sz w:val="24"/><w:szCs w:val="24"/>'
            hs = '<w:webHidden/><w:noProof/><w:sz w:val="24"/><w:szCs w:val="24"/>'
        else:
            ppr = ('<w:pPr><w:pStyle w:val="TOC2"/><w:tabs><w:tab w:val="right" w:leader="dot" '
                   'w:pos="8299"/></w:tabs><w:spacing w:line="360" w:lineRule="auto"/>'
                   '<w:ind w:left="400"/><w:rPr><w:rFonts w:eastAsia="宋体"/><w:noProof/>'
                   '<w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr></w:pPr>')
            rs = ('<w:rStyle w:val="a5"/><w:rFonts w:eastAsia="宋体" w:cs="宋体"/><w:bCs/>'
                  '<w:noProof/><w:sz w:val="24"/><w:szCs w:val="24"/>')
            hs = ('<w:rFonts w:eastAsia="宋体"/><w:webHidden/><w:noProof/><w:sz w:val="24"/>'
                  '<w:szCs w:val="24"/>')
        fld = ''
        if first:
            fld = ('<w:r><w:rPr><w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr>'
                   '<w:fldChar w:fldCharType="begin"/></w:r>'
                   '<w:r><w:rPr><w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr>'
                   '<w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText></w:r>'
                   '<w:r><w:rPr><w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr>'
                   '<w:fldChar w:fldCharType="separate"/></w:r>')
        parts.append(
            '<w:p %s>%s%s<w:hyperlink w:anchor="%s" w:history="1">'
            '<w:r><w:rPr>%s</w:rPr><w:t xml:space="preserve">%s</w:t></w:r>'
            '<w:r><w:rPr>%s</w:rPr><w:tab/></w:r>'
            '<w:r><w:rPr>%s</w:rPr><w:fldChar w:fldCharType="begin"/></w:r>'
            '<w:r><w:rPr>%s</w:rPr><w:instrText xml:space="preserve"> PAGEREF %s \\h </w:instrText></w:r>'
            '<w:r><w:rPr>%s</w:rPr><w:fldChar w:fldCharType="separate"/></w:r>'
            '<w:r><w:rPr>%s</w:rPr><w:t>%d</w:t></w:r>'
            '<w:r><w:rPr>%s</w:rPr><w:fldChar w:fldCharType="end"/></w:r>'
            '</w:hyperlink></w:p>'
            % (NS, ppr, fld, name, rs, esc(text), hs, hs, hs, name, hs, hs, pg, hs))

    parts.append(
        '<w:p %s><w:pPr><w:pStyle w:val="11"/><w:ind w:firstLine="0"/></w:pPr>'
        '<w:r><w:rPr><w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr>'
        '<w:fldChar w:fldCharType="end"/></w:r></w:p>' % NS)

    sdt = parse_xml(
        '<w:sdt %s><w:sdtPr><w:id w:val="-1457872018"/><w:docPartObj>'
        '<w:docPartGallery w:val="Table of Contents"/><w:docPartUnique/></w:docPartObj></w:sdtPr>'
        '<w:sdtContent>%s</w:sdtContent></w:sdt>' % (NS, ''.join(parts)))
    body.insert(1, sdt)

    # 让 Word / WPS 打开时自动刷新目录页码
    st = doc.settings.element
    uf = st.find(qn('w:updateFields'))
    if uf is None:
        uf = OxmlElement('w:updateFields')
        succ = ['w:hdrShapeDefaults', 'w:footnotePr', 'w:endnotePr', 'w:compat', 'w:docVars',
                'w:rsids', 'm:mathPr', 'w:attachedSchema', 'w:themeFontLang',
                'w:clrSchemeMapping', 'w:doNotIncludeSubdocsInStats', 'w:doNotAutoCompressPictures',
                'w:forceUpgrade', 'w:captions', 'w:readModeInkLockDown', 'w:smartTagType',
                'w:shapeDefaults', 'w:decimalSymbol', 'w:listSeparator']
        anchor = None
        for tag in succ:
            found = st.find(qn(tag))
            if found is not None:
                anchor = found; break
        if anchor is not None:
            st.insert(list(st).index(anchor), uf)
        else:
            st.append(uf)
    uf.set(qn('w:val'), 'true')
